"""
Document parsers: PDF (pdfplumber), DOCX (python-docx), XLSX (openpyxl).
All three dependencies are optional — if not installed the parser is silently
skipped and the file will be reported as unsupported.
"""
from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any, Dict, List

from .base import BaseParser, ParseResult

_PREVIEW = 2_000
_MAX_XLSX_ROWS = 10_000


# ─── PDF ──────────────────────────────────────────────────────────────────────

class PdfParser(BaseParser):
    SUPPORTED_EXTENSIONS = (".pdf",)
    PARSER_NAME = "pdf"

    async def parse(self, path: Path) -> ParseResult:
        result = self._base_result(path)
        loop = asyncio.get_event_loop()
        try:
            text, meta = await loop.run_in_executor(None, _parse_pdf, path)
            result.content_preview = text[:_PREVIEW]
            result.search_text = text[:100_000]
            result.char_count = len(text)
            result.record_count = meta.get("pages", 0)
            result.metadata = meta
            result.summary = (
                f"📕 PDF document · {meta.get('pages', '?')} pages · "
                f"{len(text):,} chars extracted"
            )
        except Exception as exc:
            result.success = False
            result.error = str(exc)
        return result


def _parse_pdf(path: Path) -> tuple:
    try:
        import pdfplumber  # optional
    except ImportError as e:
        raise RuntimeError("pdfplumber not installed — PDF parsing unavailable") from e

    text_parts: List[str] = []
    pages = 0
    meta_raw: Dict[str, Any] = {}

    with pdfplumber.open(path) as pdf:
        pages = len(pdf.pages)
        if pdf.metadata:
            meta_raw = {k: str(v) for k, v in pdf.metadata.items() if v}
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)

    text = "\n".join(text_parts)
    meta: Dict[str, Any] = {
        "pages": pages,
        "title": meta_raw.get("Title", ""),
        "author": meta_raw.get("Author", ""),
        "creator": meta_raw.get("Creator", ""),
        "raw_meta": meta_raw,
    }
    return text, meta


# ─── DOCX ─────────────────────────────────────────────────────────────────────

class DocxParser(BaseParser):
    SUPPORTED_EXTENSIONS = (".docx",)
    PARSER_NAME = "docx"

    async def parse(self, path: Path) -> ParseResult:
        result = self._base_result(path)
        loop = asyncio.get_event_loop()
        try:
            text, meta = await loop.run_in_executor(None, _parse_docx, path)
            result.content_preview = text[:_PREVIEW]
            result.search_text = text[:100_000]
            result.char_count = len(text)
            result.record_count = meta.get("paragraphs", 0)
            result.metadata = meta
            result.summary = (
                f"📝 Word document · {meta.get('paragraphs', '?')} paragraphs · "
                f"{len(text):,} chars extracted"
            )
        except Exception as exc:
            result.success = False
            result.error = str(exc)
        return result


def _parse_docx(path: Path) -> tuple:
    try:
        from docx import Document  # type: ignore  # python-docx package
    except ImportError as e:
        raise RuntimeError("python-docx not installed — DOCX parsing unavailable") from e

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    tables_info: List[Dict] = []
    for i, tbl in enumerate(doc.tables[:20]):
        headers = [cell.text.strip() for cell in tbl.rows[0].cells] if tbl.rows else []
        tables_info.append({"name": f"table_{i+1}", "columns": headers, "row_count": len(tbl.rows)})
    meta: Dict[str, Any] = {
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "tables_info": tables_info,
    }
    return text, meta


# ─── XLSX ─────────────────────────────────────────────────────────────────────

class XlsxParser(BaseParser):
    SUPPORTED_EXTENSIONS = (".xlsx", ".xls", ".xlsm", ".ods")
    PARSER_NAME = "xlsx"

    async def parse(self, path: Path) -> ParseResult:
        result = self._base_result(path)
        loop = asyncio.get_event_loop()
        try:
            tables, meta, search = await loop.run_in_executor(None, _parse_xlsx, path)
            total_rows = sum(t.get("row_count", 0) for t in tables)
            result.tables = tables
            result.metadata = meta
            result.record_count = total_rows
            result.search_text = search
            result.content_preview = _xlsx_preview(tables)[:_PREVIEW]
            result.summary = (
                f"📊 Spreadsheet · {len(tables)} sheet(s) · {total_rows:,} total rows"
            )
        except Exception as exc:
            result.success = False
            result.error = str(exc)
        return result


def _parse_xlsx(path: Path) -> tuple:
    try:
        import openpyxl  # optional but already in requirements
    except ImportError as e:
        raise RuntimeError("openpyxl not installed — XLSX parsing unavailable") from e

    ext = path.suffix.lower()
    if ext == ".xls":
        # xlrd for legacy .xls (read-only, no macro execution)
        try:
            import xlrd
            wb = xlrd.open_workbook(str(path))
            return _parse_xls_xlrd(wb)
        except ImportError:
            pass  # fall through to openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    tables: List[Dict[str, Any]] = []
    search_parts: List[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)
        header = next(rows_iter, None)
        if header is None:
            continue
        cols = [str(c) if c is not None else "" for c in header]
        preview: List[Dict] = []
        row_count = 0
        for row in rows_iter:
            row_count += 1
            if row_count <= 5:
                preview.append(dict(zip(cols, [str(v) if v is not None else "" for v in row])))
            if row_count >= _MAX_XLSX_ROWS:
                break
        tables.append({
            "name": sheet_name,
            "columns": cols[:100],
            "row_count": row_count,
            "preview": preview,
        })
        search_parts.append(sheet_name)
        search_parts.extend(cols)

    wb.close()
    meta: Dict[str, Any] = {
        "sheet_count": len(tables),
        "sheet_names": [t["name"] for t in tables],
    }
    return tables, meta, " ".join(search_parts)


def _parse_xls_xlrd(wb: Any) -> tuple:
    import xlrd
    tables: List[Dict[str, Any]] = []
    search_parts: List[str] = []
    for sh in wb.sheets():
        if sh.nrows == 0:
            continue
        cols = [str(sh.cell_value(0, c)) for c in range(sh.ncols)]
        preview: List[Dict] = []
        for r in range(1, min(6, sh.nrows)):
            preview.append({cols[c]: str(sh.cell_value(r, c)) for c in range(sh.ncols)})
        tables.append({
            "name": sh.name,
            "columns": cols[:100],
            "row_count": sh.nrows - 1,
            "preview": preview,
        })
        search_parts.extend([sh.name] + cols)
    meta: Dict[str, Any] = {
        "sheet_count": len(tables),
        "sheet_names": [t["name"] for t in tables],
    }
    return tables, meta, " ".join(search_parts)


def _xlsx_preview(tables: List[Dict]) -> str:
    lines = []
    for t in tables[:5]:
        cols = ", ".join(t.get("columns", [])[:10])
        lines.append(f"[{t['name']}] {t.get('row_count', '?')} rows — {cols}")
    return "\n".join(lines)
